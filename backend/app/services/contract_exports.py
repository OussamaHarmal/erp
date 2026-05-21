"""
Contract export helpers.
Generate PDF and Word files matching the Universal Invest Strategy
'Contrat de Domiciliation' design — faithful reproduction of the original.

Structure (per page):
  Page 1 – ATTESTATION DE DOMICILIATION
  Page 2 – CONTRAT DE DOMICILIATION

Layout (identical to original):
  - Provider name (bold italic blue, top-left) + services list (right-aligned, small)
  - Bold underlined centred title (22pt)
  - Body paragraphs (8pt, Cambria)
  - Article headings: "Titre 21" style — bold 8pt, left-aligned, underlined
  - Footer (centred, small): address / RC / IF / ICE
  - Signature line with tab stop

Where to place this file:
  - Drop it next to your existing contract_export.py (same location in your project).
  - Logo path: <project_root>/assets/ui_logo.jpeg  (unchanged)
"""

from datetime import datetime
from pathlib import Path

# ── python-docx ───────────────────────────────────────────────────────────────
from docx import Document as WordDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Emu

# ── ReportLab ─────────────────────────────────────────────────────────────────
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, Image, HRFlowable,
)

# ─────────────────────────────────────────────────────────────────────────────
# Paths
# ─────────────────────────────────────────────────────────────────────────────
EXPORT_DIR = Path("uploads/contracts")
LOGO_PATH  = Path(__file__).resolve().parents[1] / "assets" / "ui_logo.jpeg"

# ─────────────────────────────────────────────────────────────────────────────
# Brand constants (FIXED — our company)
# ─────────────────────────────────────────────────────────────────────────────
BLUE   = colors.HexColor("#1A3A5C")
ORANGE = colors.HexColor("#F4933A")
BORDER = colors.HexColor("#D7DEE6")

PROVIDER = {
    "name":              "UNIVERSAL INVEST STRATEGY",
    "legal_form":        "SARL AU",
    "manager_display":   "YOUSSEF BACHRA",
    "manager_signature": "BACHRA YOUSSEF",
    "manager_cin":       "BE604671",
    "address_upper":     "ANGLE RUE EL AARAR et BD LALLA ELYACOUT, IMM1, 3ème ETAGE, APPT 8",
    "address_footer":    "Angle Rue Al AARAR et Av Lalla El Yacout 3ème, imm1 Appartement 8",
    "city":              "Casablanca",
    "phone":             "+212600800747",
    "email":             "contact@ui-strategy.com",
    "rc":                "496151",
    "patente":           "34102034",
    "if_number":         "50137892",
    "cnss":              "2507310",
    "ice":               "002752348000050",
    "services": [
        "Domiciliation Juridique",
        "Centre d'Affaires",
        "Conseil Juridique\u2013 Fiscal et Comptable",
        "Tenue de Comptabilité",
        "Diagnostic des entreprises",
        "Audit",
    ],
}

DEFAULT_NATIONALITY = "Marocaine"
DEFAULT_LEGAL_FORM  = "SARL AU"


# ─────────────────────────────────────────────────────────────────────────────
# Helper utilities
# ─────────────────────────────────────────────────────────────────────────────

def ensure_export_dir() -> Path:
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    return EXPORT_DIR


def value(obj, attr: str, default: str = "xxxxxxxx") -> str:
    val = getattr(obj, attr, None) if obj else None
    return str(val).strip() if val not in (None, "") else default


def fmt_date(dt, default: str = "__/__/____") -> str:
    return dt.strftime("%d/%m/%Y") if dt else default


def fmt_birth_date(p) -> str:
    return fmt_date(getattr(p, "birth_date", None), "__/__/____")


def profile(contract):
    return contract.client.profile if getattr(contract, "client", None) else None


def client_full_name(p) -> str:
    first = value(p, "first_name", "").strip()
    last  = value(p, "last_name",  "").strip()
    return f"{first} {last}".strip() or "xxxxxxxx"


def company_name(p) -> str:
    return value(p, "company_name", "xxxxxxxx")


def money_per_month(contract) -> str:
    if getattr(contract, "duration_months", None) and getattr(contract, "price", None):
        monthly = float(contract.price) / int(contract.duration_months)
        return f"{int(monthly)}/mois" if monthly == int(monthly) else f"{monthly:.2f}/mois"
    if getattr(contract, "price", None):
        return f"{float(contract.price):.2f} MAD"
    return "167/mois"


def build_context(contract) -> dict:
    """Build variable client/contract data — all values come from DB/form."""
    p = profile(contract)
    return {
        "company_name":               company_name(p),
        "company_legal_form":         DEFAULT_LEGAL_FORM,
        "company_ice":                value(p, "company_ice"),
        "company_rc":                 value(p, "company_rc"),
        "company_address":            value(p, "company_address", value(p, "address")),
        "company_activity":           value(p, "company_activity"),
        "company_email":              value(p, "company_email", value(p, "email")),
        "company_phone":              value(p, "company_phone", value(p, "phone")),
        "representative_name":        client_full_name(p),
        "representative_nationality": DEFAULT_NATIONALITY,
        "representative_birth_date":  fmt_birth_date(p),
        "representative_cin":         value(p, "cin_number"),
        "representative_phone":       value(p, "phone"),
        "representative_email":       value(p, "company_email", value(p, "email")),
        "representative_address":     value(p, "address"),
        "start_date":                 fmt_date(getattr(contract, "start_date", None)),
        "end_date":                   fmt_date(getattr(contract, "end_date",   None)),
        "today":                      fmt_date(datetime.utcnow()),
        "monthly_fee":                money_per_month(contract),
        "contract_number":            value(contract, "contract_number", ""),
    }


# ─────────────────────────────────────────────────────────────────────────────
# Text block definitions
# Each entry: (text, kind)
# kind values: "title" | "heading" | "normal" | "bold_normal" |
#              "signature" | "signature_line" | "footer"
# ─────────────────────────────────────────────────────────────────────────────

def attestation_blocks(contract):
    c = build_context(contract)
    P = PROVIDER
    return [
        ("ATTESTATION DE DOMICILIATION", "title"),
        (
            f"Nous soussignés, <<{P['name']}>> {P['legal_form']}, "
            "attestons par la présente que :",
            "normal",
        ),
        (
            f"La société «{c['company_name']}» {c['company_legal_form']} "
            "a domicilié son adresse fiscale dans nos locaux situés à :",
            "normal",
        ),
        (
            f"{P['address_upper']} pour une période allant du    "
            f"{c['start_date']} au {c['end_date']}",
            "bold_normal",
        ),
        (
            "Nous déclarons en outre avoir pris connaissance qu'en application "
            "des dispositions de l'article 93 du CRCP, les rôles d'impôts, états "
            "de produits et autres titres de perception régulièrement émis sont "
            "exécutoires contre les redevables qui y sont inscrits, toutes "
            "personnes auprès desquelles les redevables ont élu domicile fiscal, "
            "avec leur accord.",
            "normal",
        ),
        (
            "Les personnes auprès desquelles les redevables ont élu domicile "
            "fiscal avec accord, peuvent, de ce fait, faire l'objet d'action en "
            "recouvrement au même titre que les redevables à raison de la créance "
            "due au titre de l'activité concernée par la domiciliation.",
            "normal",
        ),
        (
            "En foi de quoi, la présente attestation est délivrée pour lui "
            "permettre de procéder aux formalités administratives",
            "normal",
        ),
        (f"Fait à Casablanca Le :  {c['today']}", "bold_normal"),
        (P["manager_signature"], "signature"),
    ]


def contract_blocks(contract):
    c = build_context(contract)
    P = PROVIDER
    return [
        ("CONTRAT DE DOMICILIATION", "title"),
        (
            f"A/ Le cabinet '' {P['name']} '' {P['legal_form']}, "
            f"représenté par son gérant-unique M.{P['manager_display']} "
            f"titulaire de la CIN N° {P['manager_cin']}, "
            f"ci-après dénommé '{P['name']} ','d'une part, et d'autre part :",
            "normal",
        ),
        (
            f"La société « {c['company_name']} » {c['company_legal_form']} "
            "représenté par :",
            "normal",
        ),
        (
            f"Mr {c['representative_name']}  de nationalité "
            f"{c['representative_nationality']} , né le"
            f"{c['representative_birth_date']}, "
            f"titulaire de CIN n°{c['representative_cin']},"
            f"Téléphone N°: {c['representative_phone']}, "
            f"adress Email:{c['representative_email']}, "
            f"demeurant à {c['representative_address']}",
            "normal",
        ),
        (
            "La présente domiciliation est établie dans le cadre de la loi "
            "marocaine, notamment les mesures engagées pour faciliter "
            "l'investissement de la création d'entreprise par les jeunes "
            "promoteurs. Elle est aussi régie par le code des obligations et "
            "contrats ainsi que par les documents annexes à la présente "
            "domiciliation.",
            "normal",
        ),
        ("ARTICLEII-OBJET", "heading"),
        (
            f"Par le présent engagement de domiciliation, le cabinet "
            f"{P['name']} {P['legal_form']} s'engage moyennent une rétribution "
            "Mensuelle, à mettre a la disposition du DOMICILIE qui accepte pour "
            "la durée et aux conditions fixées par la loi marocaine et par les "
            "conditions particulières et Générales, établies dans le présent "
            "engagement de domiciliation, un ensemble de prestation tel que "
            "défini si après:",
            "normal",
        ),
        (
            "La domiciliation de son entreprise (siège social et adresse "
            "commercial); La réception, la tutelle et la mise à disposition du "
            "courrier reçu; Réception des télécopies",
            "normal",
        ),
        ("ARTICLESIII-DUREE", "heading"),
        (
            f"La présente domiciliation commence à courir Du {c['start_date']} "
            f"au {c['end_date']} Elle sera résilié automatiquement sans préavis "
            f"ou un écrit fait par {P['name']}",
            "normal",
        ),
        (
            f"Et {P['name']} n'est plus responsable des préjudices générés par "
            "le client.",
            "normal",
        ),
        (
            f"Elle sera renouvelée dans le cas où le client fait une demande "
            f"écrite acceptée par {P['name']}",
            "normal",
        ),
        ("ARTICLEVI\u2013OBLIGATION DU DOMICILIE", "heading"),
        (
            f"Le DOMICILIE s'engage à régler, aux échéances de renouvellement, "
            "les redevances relatives au frais de domiciliation ainsi que tous "
            f"les frais annexe facturés soit {c['monthly_fee']}",
            "normal",
        ),
        (
            "Le DOMICILIE s'engage à déclarer sans délai au domiciliataire "
            "selon les cas, soit tout changement relatif a son domicile "
            "personnel, soit s'il s'agit d'une personne morale, tout changement "
            "relatif a sa forme juridique, son objet, ainsi qu'au nom et au "
            "domicile personnel des personnes ayant le pouvoir générale de "
            "l'engager",
            "normal",
        ),
        (
            f"En cas de non-respect des présentes, '{P['name']}' SARL AU pourra "
            "unilatéralement et à tout moment révoquer sans formalité ni "
            "indemnité le présent engagement de domiciliation, ses obligations "
            "seront alors suspendues sans contrepartie de plein droit et sa "
            "responsabilité dégagée",
            "normal",
        ),
        (
            "Le DOMICILIE s'oblige à remettre annuellement à la société "
            "domiciliataire les copies des reçus attestant du dépôt des "
            "différentes déclarations fiscales exigibles par la loi marocaine, "
            "notamment le bilan annuel et les déclarations de TVA.",
            "normal",
        ),
        ("ARTICLEV\u2013RESILIATION DU CONTRAT", "heading"),
        (
            f"Le présent contrat pourra être résilié de plein droit par le "
            f"cabinet ''{P['name']}'' SARL AU, 30 jours après l'envoi au",
            "normal",
        ),
        (
            "DOMICILIE d'une mise en demeure par lettre recommandée avec avis "
            "de réception, restée sans effet dans les cas suivant: Non "
            "observation par le DOMICILIE de l'une quelconque des dispositions "
            "du présent engagement;",
            "normal",
        ),
        (
            "Non-paiement à leur échéance, des honoraires et/ou prestation de "
            "service; Défaut de dépôt de la déclaration fiscale légale;",
            "normal",
        ),
        (
            f"Défaut d'information du cabinet '{P['name']}'' SARL AU d'un "
            "éventuel changement dans sa situation.",
            "normal",
        ),
        ("ARTICLEVI\u2013ELECTION DE DOMICILE", "heading"),
        (
            "Pour l'exécution des présentes, les parties font élections de "
            "domicile chacune à son adresse portée sur le présent contrat et "
            "pour le DOMICILIE à son adresse personnelle ou à celle de son "
            f"représentant légal. Tout changement d'adresse du DOMICILIE n'est "
            f"opposable au cabinet '' {P['name']}' '{P['legal_form']} que s'il "
            "lui a été notifié par le DOMICILIE par lettre recommandée avec "
            "accusé de réception",
            "normal",
        ),
        (
            f"{P['name']} informe les autorités comptantes l'administration des "
            "impôts, la trésorerie  générale et de l'administration de la  "
            "Douane, le cas échéant, dans un délai n'excédant pas 15 jours  de "
            "la date de réception des plis recommandés par les services fiscaux "
            "qui n'auront pas été remis aux personnes domiciliées.",
            "normal",
        ),
        (
            "N.B : Cette attestation de domiciliation est délivrée pour la "
            "création d'une nouvelle société et n'est pas valable pour un "
            "transfert de siège social",
            "normal",
        ),
        ("ARTICLE VII\u2013FRAIS", "heading"),
        ("Les frais de légalisation sont Supportés par le domicilié.", "normal"),
        (
            f"Mr. {c['representative_name']}"
            f"\t         Mr. {P['manager_display']}",
            "signature_line",
        ),
    ]


# ═══════════════════════════════════════════════════════════════════════════════
#  PDF GENERATOR
# ═══════════════════════════════════════════════════════════════════════════════

def _pdf_styles():
    """
    ReportLab built-in fonts: Times-Roman / Times-Bold / Times-BoldItalic.
    These are the closest standard matches to the original Cambria serif font.
    """
    base = getSampleStyleSheet()
    return {
        "provider_name": ParagraphStyle(
            "provider_name", parent=base["Normal"],
            fontName="Times-BoldItalic", fontSize=11,
            textColor=BLUE, leading=14,
            spaceAfter=1,
        ),
        "provider_service": ParagraphStyle(
            "provider_service", parent=base["Normal"],
            fontName="Times-Roman", fontSize=8, leading=10,
            alignment=2,   # RIGHT
        ),
        "title": ParagraphStyle(
            "title", parent=base["Normal"],
            fontName="Times-Bold", fontSize=14,
            alignment=1,
            spaceBefore=6, spaceAfter=8, leading=17,
        ),
        "heading": ParagraphStyle(
            "heading", parent=base["Normal"],
            fontName="Times-Bold", fontSize=8,
            spaceBefore=6, spaceAfter=3, leading=11,
        ),
        "normal": ParagraphStyle(
            "normal", parent=base["Normal"],
            fontName="Times-Roman", fontSize=8, leading=11, spaceAfter=4,
        ),
        "bold_normal": ParagraphStyle(
            "bold_normal", parent=base["Normal"],
            fontName="Times-Bold", fontSize=8, leading=11, spaceAfter=4,
        ),
        "signature": ParagraphStyle(
            "signature", parent=base["Normal"],
            fontName="Times-Bold", fontSize=8,
            alignment=1,
            leading=11, spaceBefore=14,
        ),
        "signature_line": ParagraphStyle(
            "signature_line", parent=base["Normal"],
            fontName="Times-Bold", fontSize=8,
            leading=11, spaceBefore=16,
        ),
        "footer": ParagraphStyle(
            "footer", parent=base["Normal"],
            fontName="Times-Roman", fontSize=7,
            alignment=1, leading=10,
        ),
    }


def generate_contract_pdf(contract) -> str:
    """Generate a two-page PDF (attestation + contract) matching the original."""
    export_dir = ensure_export_dir()
    path = export_dir / f"{contract.contract_number}.pdf"

    PAGE_W, PAGE_H = A4
    L_MARGIN = R_MARGIN = 2.29 * cm
    T_MARGIN = 2.33 * cm
    B_MARGIN = 0.49 * cm
    CONTENT_W = PAGE_W - L_MARGIN - R_MARGIN

    doc = SimpleDocTemplate(
        str(path), pagesize=A4,
        leftMargin=L_MARGIN, rightMargin=R_MARGIN,
        topMargin=T_MARGIN, bottomMargin=B_MARGIN,
    )

    ST = _pdf_styles()
    story = []

    def _build_page(blocks_fn):
        # ── Header row: provider name (left) + services (right) ──────────────
        logo_cell = []
        if LOGO_PATH.exists():
            logo_cell.append(Image(str(LOGO_PATH), width=3.5 * cm, height=2.2 * cm))

        # Provider name + services stacked in left cell
        left_items = [Paragraph(f"<b><i>{PROVIDER['name']}</i></b>", ST["provider_name"])]
        for svc in PROVIDER["services"]:
            left_items.append(Paragraph(f"-{svc}", ST["normal"]))

        # Right cell: services right-aligned (original has them right-aligned)
        right_items = []
        for svc in PROVIDER["services"]:
            right_items.append(Paragraph(f"-{svc}", ST["provider_service"]))

        # Two-column header: [logo + name | services]
        hdr = Table(
            [[left_items, right_items]],
            colWidths=[CONTENT_W * 0.55, CONTENT_W * 0.45],
        )
        hdr.setStyle(TableStyle([
            ("VALIGN",        (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING",   (0, 0), (-1, -1), 0),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 0),
            ("TOPPADDING",    (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        story.append(hdr)
        story.append(Spacer(1, 4))

        # ── Content blocks ────────────────────────────────────────────────────
        for text, kind in blocks_fn(contract):
            if kind == "heading":
                story.append(HRFlowable(
                    width="100%", thickness=0.5, color=ORANGE, spaceAfter=2))
            story.append(Paragraph(text, ST[kind]))

        # ── Footer bar ────────────────────────────────────────────────────────
        story.append(Spacer(1, 10))
        story.append(HRFlowable(width="100%", thickness=0.4, color=BORDER))
        story.append(Spacer(1, 2))
        P = PROVIDER
        story.append(Paragraph(
            f"{P['address_footer']}  Tél:{P['phone']}", ST["footer"]))
        story.append(Paragraph(
            f"Email:{P['email']}", ST["footer"]))
        story.append(Paragraph(
            f"RC : {P['rc']} – patente : {P['patente']} – I.F : {P['if_number']}",
            ST["footer"]))
        story.append(Paragraph(
            f"–CNSS : {P['cnss']}  ICE:{P['ice']}", ST["footer"]))

    _build_page(attestation_blocks)
    story.append(PageBreak())
    _build_page(contract_blocks)

    doc.build(story)
    return str(path)


# ═══════════════════════════════════════════════════════════════════════════════
#  WORD (DOCX) GENERATOR
# ═══════════════════════════════════════════════════════════════════════════════
# Reproduces the exact paragraph styles found in the original .docx:
#   - Provider name  → "Title" style (bold italic blue, 24pt, centred with indent)
#   - Services list  → Normal, 8pt, right-aligned
#   - Page title     → Normal, 22pt bold underlined, centred
#   - Article heads  → "Titre 21" style (8pt bold, underlined)
#   - Body text      → "Body Text" style (8pt Cambria)
#   - Bold body      → Normal 8pt bold
#   - Signature/date → Normal 8pt bold, justified
#   - Sig two-col    → Normal 8pt bold, tab stop at ~4700 twips
#   - Footer lines   → Normal 8pt centred

BLUE_RGB   = RGBColor(0x1A, 0x3A, 0x5C)
ORANGE_HEX = "F4933A"


def _xml_el(tag: str) -> OxmlElement:
    return OxmlElement(tag)


def _add_bottom_border(para, color_hex: str = ORANGE_HEX):
    """Add an orange bottom border to paragraph (Titre 21 style).
    The pBdr element must come before jc and rPr in the schema order.
    """
    pPr = para._p.get_or_add_pPr()
    pBdr = _xml_el("w:pBdr")
    bottom = _xml_el("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    # Insert pBdr at position 0 (before spacing/ind/jc/rPr)
    pPr.insert(0, pBdr)


def _run(para, text: str, bold=False, size_pt=8, italic=False,
         underline=False, color: RGBColor = None, font_name: str = "Cambria"):
    """Add a styled run to a paragraph."""
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if underline:
        run.underline = True
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def _para(doc, alignment=None, space_before_pt=0, space_after_pt=3,
          line_spacing_pt=None, keep_with_next=False):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before_pt)
    fmt.space_after  = Pt(space_after_pt)
    if line_spacing_pt:
        from docx.shared import Pt as _Pt
        fmt.line_spacing = _Pt(line_spacing_pt)
    return p


def _add_tab_stop(para, pos_twips: int):
    """Add a left tab stop to a paragraph.
    The tabs element must be inserted before jc/rPr for schema compliance.
    """
    pPr  = para._p.get_or_add_pPr()
    tabs = _xml_el("w:tabs")
    tab  = _xml_el("w:tab")
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab)
    pPr.insert(0, tabs)


def _word_add_block(doc, text: str, kind: str):
    """Render one content block as a Word paragraph, matching original styles."""

    if kind == "title":
        # Original: Normal, 22pt bold underlined, centred, slight indent
        p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before_pt=1, space_after_pt=6)
        _run(p, text, bold=True, size_pt=22, underline=True)

    elif kind == "heading":
        # Original: Titre 21 — 8pt bold, left, with orange bottom border
        p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before_pt=4, space_after_pt=2)
        _run(p, text, bold=True, size_pt=8, underline=True)
        _add_bottom_border(p, ORANGE_HEX)

    elif kind == "normal":
        p = _para(doc, space_after_pt=3)
        _run(p, text, size_pt=8)

    elif kind == "bold_normal":
        p = _para(doc, space_after_pt=3)
        _run(p, text, bold=True, size_pt=8)

    elif kind == "signature":
        # Single bold name right-justified (manager signature on attestation page)
        # Original has "BACHRA YOUSSEF" bold, right-side of page
        p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                  space_before_pt=10, space_after_pt=3)
        _run(p, text, bold=True, size_pt=8)

    elif kind == "signature_line":
        # Two-column signature line: "Mr. CLIENT   \t   Mr. MANAGER"
        # Original uses a tab stop at ~6766 twips
        p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before_pt=12, space_after_pt=3)
        _add_tab_stop(p, 6766)
        # Split text on \t to apply bold formatting
        parts = text.split("\t")
        for i, part in enumerate(parts):
            if i > 0:
                p.add_run("\t")
            _run(p, part, bold=True, size_pt=8)

    return p


def _word_add_provider_header(doc):
    """Add the provider name + services header matching the original layout.
    
    Original layout:
      - Provider name: centered, bold italic, ~24pt (Title style)
      - Services: right-aligned, 8pt, each on its own line
    """
    P = PROVIDER

    # Provider name — centered bold italic, large (matching "Title" style in original)
    p_name = _para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   space_before_pt=0, space_after_pt=4)
    _run(p_name, P["name"], bold=True, italic=True, size_pt=24, color=BLUE_RGB)

    # Services — right-aligned, 8pt (matching original)
    for svc in P["services"]:
        ps = _para(doc, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                   space_before_pt=0, space_after_pt=0)
        _run(ps, f"-{svc}", size_pt=8)


def _word_build_footer(section):
    """
    Write the legal footer lines into the real Word section footer,
    so they appear anchored at the very bottom of every page — exactly
    like the original document.

    Original footer (centred, 8pt):
      Angle Rue Al AARAR et Av Lalla El Yacout 3ème, imm1 Appartement 8  Tél:+212600800747
      Email:contact@ui-strategy.com
      RC : 496151 – patente : 34102034 – I.F : 50137892
      –CNSS : 2507310  ICE:002752348000050
    """
    P = PROVIDER
    lines = [
        f"{P['address_footer']}  Tél:{P['phone']}",
        f"Email:{P['email']}",
        f"RC : {P['rc']} – patente : {P['patente']} – I.F : {P['if_number']}",
        f"–CNSS : {P['cnss']}  ICE:{P['ice']}",
    ]

    footer = section.footer

    # Force-create the footer XML definition so it is NOT linked to previous section.
    # Without this, python-docx leaves the footer "linked" (empty) for section 0.
    if not footer._has_definition:
        footer._add_definition()

    existing_paras = footer.paragraphs
    for i, line in enumerate(lines):
        if i < len(existing_paras):
            p = existing_paras[i]
            # Clear any existing runs
            for run in list(p.runs):
                run._element.getparent().remove(run._element)
        else:
            p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        _run(p, line, size_pt=8)


def generate_contract_word(contract) -> str:
    """
    Generate a two-page Word document (attestation + contract).
    Each page has the legal footer anchored at the very bottom via
    Word's section footer mechanism.
    """
    export_dir = ensure_export_dir()
    path = export_dir / f"{contract.contract_number}.docx"
    doc = WordDocument()

    # ── Margin constants (from the original document) ────────────────────────
    PAGE_W   = Cm(21.01)
    PAGE_H   = Cm(29.70)
    TOP      = Cm(2.33)
    BOTTOM   = Cm(1.80)   # enough room for 4 footer lines at 8pt
    LEFT     = Cm(2.29)
    RIGHT    = Cm(2.29)
    FOOTER_D = Cm(0.49)   # footer distance from bottom edge (original bottom margin)

    from docx.oxml.ns import qn as _qn
    from docx.oxml   import OxmlElement as _el

    def _apply_section_margins(section):
        section.page_width        = PAGE_W
        section.page_height       = PAGE_H
        section.top_margin        = TOP
        section.bottom_margin     = BOTTOM
        section.left_margin       = LEFT
        section.right_margin      = RIGHT
        section.footer_distance   = FOOTER_D

    # ── Page 1 body: Attestation ─────────────────────────────────────────────
    _word_add_provider_header(doc)
    doc.add_paragraph()   # blank spacer

    for text, kind in attestation_blocks(contract):
        _word_add_block(doc, text, kind)

    # ── Section break: end of page 1, start of page 2 ───────────────────────
    # Embed a sectPr in the last paragraph of page 1 to create section 1.
    # This sectPr carries the page 1 margin/size settings.
    p_brk = doc.add_paragraph()
    pPr   = p_brk._p.get_or_add_pPr()
    sectPr = _el("w:sectPr")
    pgSz   = _el("w:pgSz")
    pgSz.set(_qn("w:w"), str(int(PAGE_W.twips)))
    pgSz.set(_qn("w:h"), str(int(PAGE_H.twips)))
    pgMar  = _el("w:pgMar")
    pgMar.set(_qn("w:top"),    str(int(TOP.twips)))
    pgMar.set(_qn("w:right"),  str(int(RIGHT.twips)))
    pgMar.set(_qn("w:bottom"), str(int(BOTTOM.twips)))
    pgMar.set(_qn("w:left"),   str(int(LEFT.twips)))
    pgMar.set(_qn("w:footer"), str(int(FOOTER_D.twips)))
    pgMar.set(_qn("w:gutter"), "0")
    sectPr.append(pgSz)
    sectPr.append(pgMar)
    pPr.append(sectPr)

    # ── Page 2 body: Contract ────────────────────────────────────────────────
    _word_add_provider_header(doc)
    doc.add_paragraph()   # blank spacer

    for text, kind in contract_blocks(contract):
        _word_add_block(doc, text, kind)

    # ── Apply margins to both sections and write footers ─────────────────────
    # IMPORTANT: footers must be written AFTER all body content is added,
    # otherwise python-docx may overwrite the section 0 sectPr definition.
    _apply_section_margins(doc.sections[0])
    _word_build_footer(doc.sections[0])

    _apply_section_margins(doc.sections[-1])
    _word_build_footer(doc.sections[-1])

    doc.save(str(path))
    return str(path)
